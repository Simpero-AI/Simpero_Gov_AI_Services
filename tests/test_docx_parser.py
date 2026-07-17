"""DOCX path tests.

Same convention as the rest of the parser suite: fast, CI-portable tests against
synthetic documents generated in-process (python-docx here, as reportlab does for
PDF and openpyxl for XLSX) — no committed binary fixtures.

The point of these tests is that the DOCX lane clears the SAME trust bar as the
PDF lane through the same exact-span core, while carrying the provenance shape a
flow document actually has: (paragraph, char_start, char_end) and no geometry.
"""

from io import BytesIO

import pytest
from docx import Document
from fastapi.testclient import TestClient

from parser_service.docx_parser import parse_docx_bytes
from parser_service.errors import ParseError
from parser_service.main import app
from parser_service.resolver import resolve_in_paragraph
from parser_service.schemas import ParagraphIndex


def _docx_bytes(paragraphs: list[str]) -> bytes:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# parse_docx_bytes -- the paragraph index.
# --------------------------------------------------------------------------- #


def test_paragraph_index_carries_text_in_body_order() -> None:
    result = parse_docx_bytes(_docx_bytes(["First paragraph.", "Second paragraph."]))

    assert [p.paragraph for p in result.paragraphs] == [0, 1]
    assert result.paragraphs[0].text == "First paragraph."
    assert result.paragraphs[1].text == "Second paragraph."
    assert len(result.sha256) == 64


def test_empty_paragraphs_are_skipped_not_indexed() -> None:
    # A blank paragraph is not citable; indexing it would shift every subsequent
    # locator for no benefit.
    result = parse_docx_bytes(_docx_bytes(["Real content.", "   ", "", "More content."]))

    assert [p.text for p in result.paragraphs] == ["Real content.", "More content."]
    # Locators stay contiguous over the paragraphs that actually exist.
    assert [p.paragraph for p in result.paragraphs] == [0, 1]


def test_numeric_tokens_normalized_like_the_page_index() -> None:
    # The same rule as DS-W3-1, so a quote resolves identically whichever lane
    # read the document -- the extractor never has to know the source format.
    result = parse_docx_bytes(_docx_bytes(["Gross Margin was 3 ,817 thousand."]))

    assert "3,817" in result.paragraphs[0].text


def test_no_geometry_is_produced_for_docx() -> None:
    # A Word file has no page layout, so there is nothing to report. The absence
    # is the correct answer, not a gap: ParagraphIndex has no char_map at all.
    result = parse_docx_bytes(_docx_bytes(["Some prose."]))

    assert isinstance(result.paragraphs[0], ParagraphIndex)
    assert not hasattr(result.paragraphs[0], "char_map")


# --------------------------------------------------------------------------- #
# Fail-closed guards -- the DOCX lane rejects what the PDF lane rejects.
# --------------------------------------------------------------------------- #


def test_zero_byte_docx_rejected() -> None:
    with pytest.raises(ParseError) as exc_info:
        parse_docx_bytes(b"")
    assert exc_info.value.code == "zero_byte_docx"
    assert exc_info.value.status_code == 400


def test_corrupt_docx_rejected() -> None:
    with pytest.raises(ParseError) as exc_info:
        parse_docx_bytes(b"this is not a real docx file")
    assert exc_info.value.code == "corrupt_docx"
    assert exc_info.value.status_code == 400


def test_duplicate_sha256_rejected() -> None:
    data = _docx_bytes(["Anything."])
    first = parse_docx_bytes(data)

    with pytest.raises(ParseError) as exc_info:
        parse_docx_bytes(data, known_sha256s={first.sha256})
    assert exc_info.value.code == "duplicate_docx"
    assert exc_info.value.status_code == 409


def test_docx_with_no_extractable_text_rejected() -> None:
    # A document that yields nothing is a failure to surface, never an empty
    # success -- same rule as the PDF lane's no_extractable_text guard.
    with pytest.raises(ParseError) as exc_info:
        parse_docx_bytes(_docx_bytes(["", "   "]))
    assert exc_info.value.code == "no_extractable_text"


def test_oversized_uncompressed_docx_is_rejected(monkeypatch: pytest.MonkeyPatch) -> None:
    # The shared OOXML decompression-bomb guard applies to DOCX exactly as it
    # does to XLSX -- both are a zip of XML.
    from parser_service import config

    config.get_settings.cache_clear()
    monkeypatch.setenv("PARSER_MAX_OOXML_UNCOMPRESSED_BYTES", "100")
    config.get_settings.cache_clear()
    try:
        with pytest.raises(ParseError) as exc_info:
            parse_docx_bytes(_docx_bytes(["Any document expands past 100 bytes."]))
        assert exc_info.value.code == "docx_too_large"
        assert exc_info.value.status_code == 413
    finally:
        config.get_settings.cache_clear()


# --------------------------------------------------------------------------- #
# resolve_in_paragraph -- the same trust boundary, without geometry.
# --------------------------------------------------------------------------- #


def test_quote_resolves_to_an_exact_paragraph_span() -> None:
    result = parse_docx_bytes(_docx_bytes(["Intro.", "Revenue was $15,295 last year."]))
    target = result.paragraphs[1]

    span = resolve_in_paragraph("$15,295", target)

    assert span is not None
    assert span.paragraph == 1
    # char_end is exclusive: the slice reproduces the quote exactly.
    assert target.text[span.char_start : span.char_end] == "$15,295"


def test_absent_quote_fails_closed() -> None:
    result = parse_docx_bytes(_docx_bytes(["Revenue was $15,295 last year."]))
    assert resolve_in_paragraph("not-in-this-paragraph", result.paragraphs[0]) is None


def test_ambiguous_quote_fails_closed() -> None:
    # Found twice -> which instance the claim cites is unknowable. Same rule as
    # the PDF lane: refuse rather than guess.
    result = parse_docx_bytes(_docx_bytes(["$15,295 here and $15,295 again."]))
    assert resolve_in_paragraph("$15,295", result.paragraphs[0]) is None


def test_empty_quote_fails_closed() -> None:
    result = parse_docx_bytes(_docx_bytes(["Anything at all."]))
    assert resolve_in_paragraph("", result.paragraphs[0]) is None
    assert resolve_in_paragraph("   ", result.paragraphs[0]) is None


def test_quote_wrapping_whitespace_still_resolves() -> None:
    # Whitespace-flexible, exactly like the PDF lane -- the one principled
    # relaxation, shared via find_exact_span.
    result = parse_docx_bytes(_docx_bytes(["Total debt   to net worth ratio."]))

    span = resolve_in_paragraph("Total debt to net worth", result.paragraphs[0])

    assert span is not None
    assert result.paragraphs[0].text[span.char_start : span.char_end] == "Total debt   to net worth"


def test_resolution_is_exact_not_fuzzy() -> None:
    result = parse_docx_bytes(_docx_bytes(["Revenue was $15,295 last year."]))
    para = result.paragraphs[0]

    assert resolve_in_paragraph("15295", para) is None  # comma dropped
    assert resolve_in_paragraph("$15,295.00", para) is None  # extra characters


# --------------------------------------------------------------------------- #
# /parse dispatch -- one endpoint, format detected from the bytes.
# --------------------------------------------------------------------------- #

client = TestClient(app)


def test_parse_endpoint_dispatches_docx_by_content() -> None:
    # No Content-Type is sent on purpose: the lane is chosen from the bytes, so
    # the caller does not have to declare (or be trusted about) the format.
    data = _docx_bytes(["Revenue was $15,295 last year."])

    response = client.post("/parse", content=data)

    assert response.status_code == 200
    body = response.json()
    assert body["kind"] == "docx"
    assert body["pages"] is None and body["sheets"] is None
    assert body["paragraphs"][0]["text"] == "Revenue was $15,295 last year."
    assert response.headers["X-Content-SHA256"] == body["sha256"]


def test_parse_endpoint_rejects_duplicate_docx() -> None:
    data = _docx_bytes(["Duplicate source document."])
    digest = parse_docx_bytes(data).sha256

    response = client.post("/parse", content=data, headers={"X-Known-SHA256": digest})

    assert response.status_code == 409
    assert response.json()["detail"]["code"] == "duplicate_docx"
